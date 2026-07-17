# -*- coding: utf-8 -*-
"""
软著申请源代码文档生成器
从 Unity C# 源文件提取代码，按软著要求排版（每页50行，页眉页脚）
前30页连续开头 + 后30页连续结尾 = 60页
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ===== 配置 =====
PROJECT_SCRIPTS = r"F:\WorkBuddy\H5MiniGame\MiniGame1\MiniGame1_Project\Assets\Scripts"
PROJECT_EDITOR  = r"F:\WorkBuddy\H5MiniGame\MiniGame1\MiniGame1_Project\Assets\Editor"
OUTPUT_PATH     = r"F:\WorkBuddy\Claw\Design_Test\疯狂铲屎官_源代码文档.docx"
GAME_NAME       = "疯狂铲屎官"
VERSION         = "V1.0"
LINES_PER_PAGE  = 50
TOTAL_PAGES     = 60  # 前30页 + 后30页
FRONT_PAGES     = 30
BACK_PAGES      = 30

# 按逻辑顺序排列文件（软著审查时前面的代码应该更"核心"）
FILE_ORDER = [
    # Core
    "Core/GameEntry.cs",
    # PetGame 核心（按重要性排序）
    "PetGame/PetGameData.cs",
    "PetGame/PetLevelConfigV2.cs",
    "PetGame/PourSystem.cs",
    "PetGame/PetGameManager.cs",
    "PetGame/PetGameFSM.cs",
    "PetGame/PetGameUI.cs",
    "PetGame/LevelGenerator.cs",
    "PetGame/PetGameSpriteConfig.cs",
    # Editor 工具
    "Editor/BatchLevelGenerator.cs",
    "Editor/LevelEditorWindow.cs",
    "Editor/PetGamePrefabGenV2.cs",
    "Editor/PetGameSceneSetup.cs",
    # 废弃但保留供代码量
    "Core/GameManager.cs",
    "PetGame/PetGameSystems.cs",
]

def collect_all_code():
    """按 FILE_ORDER 顺序读取所有源文件代码"""
    all_lines = []
    scripts_root = Path(PROJECT_SCRIPTS)
    editor_root = Path(PROJECT_EDITOR)

    for rel_path in FILE_ORDER:
        if rel_path.startswith("Editor/"):
            full_path = editor_root / Path(rel_path).name
        else:
            full_path = scripts_root / rel_path
        if not full_path.exists():
            print(f"[WARN] 文件不存在: {full_path}")
            continue
        with open(full_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        # 添加文件分隔注释
        all_lines.append(f"// ===== 文件: {Path(rel_path).name} =====\n")
        for line in lines:
            line = line.rstrip('\n').rstrip('\r')
            all_lines.append(line + "\n")
        all_lines.append("\n")  # 文件间空行
        print(f"  {rel_path}: {len(lines)} 行")

    return all_lines


def create_document():
    doc = Document()

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10.5)
    # 设置中文字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 行间距固定值，确保每页50行
    para_fmt = style.paragraph_format
    para_fmt.line_spacing = Cm(0.52)  # 精确控制每页行数
    para_fmt.space_before = Pt(0)
    para_fmt.space_after = Pt(0)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{GAME_NAME} {VERSION} 源程序"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.name = 'Courier New'

    # 页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加页码字段
    run = footer_para.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)

    run2 = footer_para.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = footer_para.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)

    for r in footer_para.runs:
        r.font.size = Pt(9)

    return doc


def add_code_page(doc, lines, page_num):
    """添加一页代码（50行）"""
    for i, line in enumerate(lines):
        # 替换 tab 为 4 空格
        line = line.replace('\t', '    ')
        # 截断过长行（避免换页）
        if len(line) > 95:
            line = line[:92] + '...'

        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = Cm(0.52)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def main():
    print("=" * 60)
    print("软著源代码文档生成器")
    print(f"游戏名: {GAME_NAME} {VERSION}")
    print(f"目标: 前{FRONT_PAGES}页 + 后{BACK_PAGES}页 = {TOTAL_PAGES}页")
    print(f"每页: {LINES_PER_PAGE}行")
    print("=" * 60)

    # 1. 收集所有代码
    print("\n[1/4] 收集源代码...")
    all_lines = collect_all_code()
    total_lines = len(all_lines)
    total_code_pages = (total_lines + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    print(f"  总行数: {total_lines}")
    print(f"  总页数: {total_code_pages}")

    if total_code_pages < TOTAL_PAGES:
        print(f"  [WARN] 代码不足{TOTAL_PAGES}页，将使用全部代码")

    # 2. 计算前30页和后30页的行范围
    front_lines = all_lines[:FRONT_PAGES * LINES_PER_PAGE]

    if total_code_pages > TOTAL_PAGES:
        back_start = total_lines - BACK_PAGES * LINES_PER_PAGE
        back_lines = all_lines[back_start:]
    else:
        back_lines = all_lines[FRONT_PAGES * LINES_PER_PAGE:]

    print(f"\n[2/4] 代码分页:")
    print(f"  前{FRONT_PAGES}页: 第1~{FRONT_PAGES * LINES_PER_PAGE}行")
    if total_code_pages > TOTAL_PAGES:
        print(f"  后{BACK_PAGES}页: 第{back_start + 1}~{total_lines}行")
    else:
        print(f"  后{BACK_PAGES}页: 第{FRONT_PAGES * LINES_PER_PAGE + 1}~{total_lines}行 (不足{BACK_PAGES}页)")

    # 3. 生成文档
    print(f"\n[3/4] 生成 Word 文档...")
    doc = create_document()

    # 前半部分
    for page in range(FRONT_PAGES):
        start = page * LINES_PER_PAGE
        end = start + LINES_PER_PAGE
        page_lines = front_lines[start:end]
        if not page_lines:
            break
        add_code_page(doc, page_lines, page + 1)
        add_page_break(doc)
        if (page + 1) % 10 == 0:
            print(f"  前{page + 1}页完成")

    # 后半部分
    total_back = (len(back_lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    for page in range(total_back):
        start = page * LINES_PER_PAGE
        end = start + LINES_PER_PAGE
        page_lines = back_lines[start:end]
        if not page_lines:
            break
        add_code_page(doc, page_lines, FRONT_PAGES + page + 1)
        if page < total_back - 1:
            add_page_break(doc)
        if (page + 1) % 10 == 0:
            print(f"  后{page + 1}页完成")

    # 4. 保存
    print(f"\n[4/4] 保存文件...")
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"  路径: {OUTPUT_PATH}")
    print(f"  大小: {file_size:.1f} KB")
    print("\n完成!")


if __name__ == '__main__':
    main()
