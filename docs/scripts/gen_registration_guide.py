# -*- coding: utf-8 -*-
"""
软著申请 — 在线申请表填写指南
中国版权保护中心在线申请系统
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_PATH = r"F:\WorkBuddy\Claw\Design_Test\疯狂铲屎官_软著申请表填写指南.docx"
GAME_NAME = "疯狂铲屎官"
VERSION = "V1.0"

def set_run_font(run, font_name='宋体', font_size=Pt(12), bold=False, color=None):
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    if level == 1:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(16), bold=True)
    elif level == 2:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(14), bold=True)
    else:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(12), bold=True)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_body(doc, text, indent=0, bold=False, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_name='宋体', font_size=Pt(12), bold=bold, color=color)
    para.paragraph_format.first_line_indent = Cm(indent)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    return para

def add_field(doc, field_name, value, note=""):
    """添加表单字段说明"""
    para = doc.add_paragraph()
    run1 = para.add_run(f"【{field_name}】")
    set_run_font(run1, font_name='宋体', font_size=Pt(12), bold=True, color=(0, 88, 168))
    
    run2 = para.add_run(f" {value}")
    set_run_font(run2, font_name='宋体', font_size=Pt(12), bold=True)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    
    if note:
        para2 = doc.add_paragraph()
        run3 = para2.add_run(f"   提示：{note}")
        set_run_font(run3, font_name='宋体', font_size=Pt(10), color=(128, 128, 128))
        para2.paragraph_format.line_spacing = 1.5
        para2.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # ===== 封面 =====
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{GAME_NAME}")
    set_run_font(run, font_name='黑体', font_size=Pt(26), bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("软件著作权登记申请表填写指南")
    set_run_font(run, font_name='黑体', font_size=Pt(22), bold=True)
    
    for _ in range(3):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("著作权人：黎婉丹\n日期：2026年07月")
    set_run_font(run, font_name='宋体', font_size=Pt(14))
    
    doc.add_page_break()
    
    # ===== 一、申请前准备 =====
    add_heading(doc, "一、申请前准备", 1)
    add_body(doc, "1. 访问中国版权保护中心官网：www.ccopyright.com.cn")
    add_body(doc, "2. 注册账号并完成实名认证（需身份证+人脸识别）")
    add_body(doc, "3. 准备好以下材料扫描件：")
    add_body(doc, "   · 身份证复印件（正反面）", indent=0.74)
    add_body(doc, "   · 软件说明书（已生成：疯狂铲屎官_软件说明书.docx）", indent=0.74)
    add_body(doc, "   · 源代码文档（已生成：疯狂铲屎官_源代码文档.docx）", indent=0.74)
    add_body(doc, "4. 将两份Word文档打印成PDF格式（在线系统上传需要PDF）")
    
    # ===== 二、申请表字段填写说明 =====
    add_heading(doc, "二、申请表字段填写说明", 1)
    add_body(doc, "登录中国版权保护中心在线系统，进入“计算机软件著作权登记”申请页面，按以下字段填写：", bold=True)
    add_body(doc, "")
    
    add_field(doc, "软件名称", "疯狂铲屎官", "必填，与说明书、源代码文档名称一致")
    add_field(doc, "版本号", "V1.0", "必填")
    add_field(doc, "软件简称", "（留空）", "如有推广用的短名称可填，否则留空")
    add_field(doc, "分类号", "选择“游戏软件”", "在分类中选择：计算机软件 > 游戏软件 > 休闲益智类")
    add_field(doc, "开发完成日期", "2026年07月13日", "填写首次完成开发的日期，可填当前日期或稍早")
    add_field(doc, "首次发表日期", "未发表", "如未公开发布，选择“未发表”")
    add_field(doc, "开发方式", "独立开发", "选择自己独立开发")
    add_field(doc, "权利归属", "个人", "选择“个人作品”")
    add_field(doc, "著作权人", "黎婉丹", "与身份证姓名完全一致")
    add_field(doc, "国籍", "中国", "自动填充或选择")
    add_field(doc, "身份证明", "居民身份证", "上传身份证扫描件")
    add_field(doc, "源程序量", "约3000行", "填写源代码总函数估算量（实际文档60页约3000行）")
    add_field(doc, "主要功能", "休闲益智游戏，通过倾倒食物完成宠物喂食", "简要描述软件功能")
    add_field(doc, "技术特点", "基于Unity引擎开发，采用FSM状态机、BFS算法、事件驱动架构", "简要描述技术特点")
    
    # ===== 三、上传材料清单 =====
    add_heading(doc, "三、上传材料清单", 1)
    add_body(doc, "在线申请时需要上传以下文件（PDF格式）：")
    add_body(doc, "1. 身份证明文件", bold=True)
    add_body(doc, "   · 身份证复印件扫描件（正反面合并为1个PDF）", indent=0.74)
    add_body(doc, "")
    add_body(doc, "2. 软件说明书", bold=True)
    add_body(doc, "   · 文件名：疯狂铲屎官_软件说明书.pdf", indent=0.74)
    add_body(doc, "   · 页数：约10-15页", indent=0.74)
    add_body(doc, "   · 要求：包含软件功能描述、界面截图、操作流程", indent=0.74)
    add_body(doc, "")
    add_body(doc, "3. 源代码文档", bold=True)
    add_body(doc, "   · 文件名：疯狂铲屎官_源代码文档.pdf", indent=0.74)
    add_body(doc, "   · 页数：60页（前30页+后30页）", indent=0.74)
    add_body(doc, "   · 要求：每页不少于50行代码，A4纸单倍行距", indent=0.74)
    add_body(doc, "   · 格式：页眉写“疯狂铲屎官 V1.0 源程序”，页脚写页码", indent=0.74)
    add_body(doc, "")
    add_body(doc, "4. 其他材料（如有）", bold=True)
    add_body(doc, "   · 如委托他人代办，需上传《代理委托书》", indent=0.74)
    add_body(doc, "   · 如合作开发，需上传《合作开发协议》", indent=0.74)
    
    # ===== 四、费用说明 =====
    add_heading(doc, "四、费用说明", 1)
    add_body(doc, "· 软件著作权登记费：免费（2022年起国家取消收费）")
    add_body(doc, "· 如需加急办理：部分省市版权中心提供加急服务，需额外付费")
    add_body(doc, "· 证书邮寄费：部分地区需自付邮寄费用")
    
    # ===== 五、办理周期 =====
    add_heading(doc, "五、办理周期", 1)
    add_body(doc, "· 普通办理：约30-60个工作日（1.5-2个月）")
    add_body(doc, "· 加急办理：约15-20个工作日（需额外付费）")
    add_body(doc, "· 证书领取：可选择邮寄或自取")
    
    # ===== 六、注意事项 =====
    add_heading(doc, "六、注意事项", 1)
    add_body(doc, "1. 材料真实性", bold=True, color=(200, 0, 0))
    add_body(doc, "   所有提交材料必须真实有效，如发现虚假信息，将影响登记结果并可能承担法律责任。")
    add_body(doc, "2. 名称一致性", bold=True, color=(200, 0, 0))
    add_body(doc, "   软件名称在申请表、说明书、源代码文档中必须完全一致，否则会被退回。")
    add_body(doc, "3. 源代码格式", bold=True, color=(200, 0, 0))
    add_body(doc, "   源代码文档必须是连续的，不能跳页、跳行。每页50行，前后各30页。")
    add_body(doc, "4. 发表状态", bold=True, color=(200, 0, 0))
    add_body(doc, "   如选择“未发表”，则保护期从登记日起算；如选择“已发表”，保护期从首次发表日起算。")
    add_body(doc, "5. 抖音上线", bold=True, color=(200, 0, 0))
    add_body(doc, "   软著申请期间可以继续开发和丰富内容，不影响后续上线。软著只是权属证明，不是上线许可。")
    add_body(doc, "6. 多平台发布", bold=True, color=(200, 0, 0))
    add_body(doc, "   软著登记的是软件本身，不限制发布平台。同一软著可用于微信、抖音、App Store等多平台。")
    
    # ===== 七、后续步骤 =====
    add_heading(doc, "七、提交后流程", 1)
    add_body(doc, "1. 提交申请 → 2. 形式审查（3-7天）→ 3. 受理通知书 → 4. 实质审查（20-30天）→ 5. 登记公告 → 6. 领取证书")
    add_body(doc, "")
    add_body(doc, "形式审查：检查材料是否齐全、格式是否正确")
    add_body(doc, "实质审查：审查软件的独创性、材料的真实性")
    add_body(doc, "登记公告：在中国版权保护中心官网公示（可查询）")
    add_body(doc, "领取证书：电子版或纸质版《计算机软件著作权登记证书》")
    
    # ===== 八、联系方式 =====
    add_heading(doc, "八、联系方式", 1)
    add_body(doc, "· 中国版权保护中心官网：www.ccopyright.com.cn")
    add_body(doc, "· 咨询电话：010-68003887（工作日9:00-17:00）")
    add_body(doc, "· 邮寄地址：北京市西城区天桥南大街1号天桥艺术大厦B座（如需纸质提交）")
    
    # 保存
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"填写指南已生成: {OUTPUT_PATH}")
    print(f"文件大小: {file_size:.1f} KB")
    print("完成!")

if __name__ == '__main__':
    main()
