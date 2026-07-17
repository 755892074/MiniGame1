from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('中国版权登记业务平台 — 软件著作权申请表填写内容', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('软件名称：疯狂铲屎官')
doc.add_paragraph('著作权人：黎婉丹')
doc.add_paragraph('填写日期：2026年7月')
doc.add_paragraph('')

# Page 1
doc.add_heading('第1页 — 软件申请信息 & 软件开发信息', level=1)

doc.add_heading('开发的硬件环境', level=2)
p = doc.add_paragraph()
p.add_run('Intel Core i5 及以上处理器，内存 8GB 及以上，硬盘 256GB 及以上').font.size = Pt(11)

doc.add_heading('运行的硬件环境', level=2)
p = doc.add_paragraph()
p.add_run('智能手机（Android 7.0+ / iOS 12.0+），需安装抖音APP').font.size = Pt(11)

doc.add_heading('开发该软件的操作系统', level=2)
p = doc.add_paragraph()
p.add_run('Windows 11').font.size = Pt(11)

doc.add_heading('软件开发环境 / 开发工具', level=2)
p = doc.add_paragraph()
p.add_run('团结引擎（Tuanjie Engine）1.9.3，Visual Studio 2022，C#').font.size = Pt(11)

# Page 2
doc.add_heading('第2页 — 软件开发信息（续）', level=1)

doc.add_heading('该软件的运行平台 / 操作系统', level=2)
p = doc.add_paragraph()
p.add_run('Android 7.0+ / iOS 12.0+（通过抖音小游戏平台运行）').font.size = Pt(11)

doc.add_heading('软件运行支撑环境 / 支持软件', level=2)
p = doc.add_paragraph()
p.add_run('抖音APP（抖音小游戏运行环境）、OpenGL ES 3.0+、WebGL 2.0+').font.size = Pt(11)

doc.add_heading('编程语言', level=2)
p = doc.add_paragraph()
p.add_run('C#').font.size = Pt(11)

doc.add_heading('源程序量', level=2)
p = doc.add_paragraph()
p.add_run('约 3155 行（15个 C# 源文件）').font.size = Pt(11)

doc.add_paragraph('')

# Page 3
doc.add_heading('第3页 — 软件功能与特点', level=1)

doc.add_heading('开发目的', level=2)
purpose = ('本软件是一款面向休闲益智类游戏爱好者的宠物主题小游戏，旨在通过简单有趣的关卡挑战为玩家带来轻松愉悦的娱乐体验。'
           '游戏融入宠物救助的公益主题，让玩家在娱乐中感受爱心传递的温暖。'
           '同时，游戏采用渐进式难度设计和丰富的成长系统，使玩家在挑战中不断提升成就感，延长游戏生命周期。')
p = doc.add_paragraph(purpose)
p.paragraph_format.line_spacing = 1.5

doc.add_heading('面向领域 / 行业', level=2)
p = doc.add_paragraph()
p.add_run('娱乐游戏 / 移动游戏 / 休闲益智').font.size = Pt(11)

doc.add_heading('软件的主要功能（500~1300字）', level=2)
features = ('本软件是一款宠物主题休闲益智类小游戏，核心玩法基于经典的三消/倒碗喂食机制。玩家通过点击碗，将碗内食物分发给对应种类的宠物，当所有宠物都吃饱后即可通关。'
            '\n\n【核心功能模块】'
            '\n1. 关卡挑战系统：包含50个精心设计的关卡，每个关卡有不同的宠物种类、碗的数量、食物种类和步数限制。'
            '关卡难度随进度逐步提升，从简单的单宠物喂食逐步扩展到多宠物、多碗、多步数限制的综合挑战。'
            '\n\n2. 宠物系统：设计了6种可爱宠物（猫咪、狗狗、兔子、仓鼠、鹦鹉、龙猫），每种宠物有不同的外观特征和喜欢的食物类型。'
            '宠物包含idle、happy、cry、run、cute五种动画状态，丰富视觉表现。'
            '\n\n3. 铲屎官成长系统：玩家拥有"铲屎官"等级称号，共8个等级（实习→初级→中级→高级→资深→专家→大师→传奇）。'
            '通过通关积累经验值提升等级，等级影响可解锁的游戏内容和特权。'
            '\n\n4. 货币系统：包含三种货币——小鱼干（通关基础奖励）、救助徽章（星级通关成就货币）、彩虹毛球（稀有货币，通过观看广告获取）。'
            '货币可用于解锁宠物、升级建筑、购买道具。'
            '\n\n5. 局外养成系统：玩家通关获得的收入用于照顾流浪宠物。宠物住所从简陋的破铁笼逐步升级为豪华别墅，反映玩家的爱心投入和成长轨迹。'
            '\n\n6. 道具与辅助功能：提供撤销上一步、添加新碗、重新洗牌三种辅助道具，帮助玩家应对困难关卡。'
            '\n\n7. 星级评价系统：每关根据剩余步数评1-3星，鼓励玩家追求更高效率和更优解。'
            '\n\n【界面与交互】'
            '\n- 主界面：关卡选择网格，显示各关通关星级'
            '\n- 游戏界面：顶部显示关卡信息、剩余步数、玩家等级状态；中部为宠物和碗的互动区域'
            '\n- 结算界面：通关后显示星级评价、获得的小鱼干/徽章/经验值奖励，以及看广告翻倍按钮')
p = doc.add_paragraph(features)
p.paragraph_format.line_spacing = 1.5

doc.add_heading('软件的技术特点', level=2)
tech = ('【技术架构】'
        '\n1. 基于团结引擎（Tuanjie Engine）1.9.3 开发，采用 Unity C# 脚本架构，利用 ECS（Entity-Component-System）组件化设计模式，'
        '代码模块化程度高，便于维护和扩展。'
        '\n\n2. 数据驱动设计：关卡配置采用 JSON 数据表驱动，通过 PetLevelConfigV2 数据模型管理50关的关卡参数（宠物种类、碗数量、食物类型、目标分数、步数限制等），'
        '便于平衡性调整和新关卡扩展。'
        '\n\n3. 状态机架构：游戏核心逻辑采用有限状态机（FSM）管理游戏流程，包括 Waiting（等待）、Dealing（发牌/分发）、Win（通关）、Fail（失败）、Pause（暂停）等状态，'
        '逻辑清晰、可维护性强。'
        '\n\n4. 关卡生成管线：采用离线关卡生成方案，在编辑器中通过 BatchLevelGenerator 工具批量生成和验证关卡，支持 BFS（广度优先搜索）死局检测，确保每个关卡均可解。'
        '\n\n5. 跨平台适配：团结引擎原生支持 WebGL 和抖音小游戏平台导出，通过条件编译实现平台差异化适配（如存档同步、广告接入等）。'
        '\n\n6. 资源管理优化：采用三级资源架构（首包<4MB、分包、CDN流式加载），控制初始包体大小。帧动画统一使用 128×128 尺寸，适配小游戏包体限制。')

# 更新技术特点精简版（100字限制）
tech_short = ('基于团结引擎1.9.3，ECS组件化架构，JSON数据驱动关卡配置，FSM状态机管理游戏流程，离线关卡生成+BFS死局检测，WebGL跨平台适配抖音小游戏，三级资源架构控制包体。')
p = doc.add_paragraph(tech)
p.paragraph_format.line_spacing = 1.5

# Page 4
doc.add_heading('第4页 — 鉴别材料', level=1)
doc.add_heading('程序鉴别材料', level=2)
p = doc.add_paragraph()
p.add_run('一般交存 — 源程序前连续的30页和后连续的30页').font.size = Pt(11)
doc.add_paragraph('请上传：疯狂铲屎官_源代码文档.pdf（已生成，共60页）')

doc.add_heading('文档鉴别材料', level=2)
p = doc.add_paragraph()
p.add_run('一般交存 — 提交任何一种文档的前连续的30页和后连续的30页').font.size = Pt(11)
doc.add_paragraph('请上传：疯狂铲屎官_软件说明书.pdf（已生成）')

doc.add_paragraph('')

# Summary
doc.add_heading('填写汇总表（对照填表用）', level=1)

table_data = [
    ['字段', '填写内容'],
    ['软件名称', '疯狂铲屎官'],
    ['软件简称', '（留空或填"铲屎官"）'],
    ['版本号', 'V1.0'],
    ['分类', '应用软件 → 游戏软件'],
    ['开发方式', '独立开发'],
    ['开发完成日期', '2026年7月13日'],
    ['著作权人', '黎婉丹'],
    ['开发的硬件环境', 'Intel Core i5及以上，内存8GB及以上，硬盘256GB及以上'],
    ['运行的硬件环境', '智能手机（Android 7.0+/iOS 12.0+），需安装抖音APP'],
    ['开发操作系统', 'Windows 11'],
    ['开发工具', '团结引擎（Tuanjie Engine）1.9.3，Visual Studio 2022，C#'],
    ['运行平台/操作系统', 'Android 7.0+ / iOS 12.0+（通过抖音小游戏平台）'],
    ['运行支撑环境', '抖音APP（抖音小游戏运行环境）、OpenGL ES 3.0+、WebGL 2.0+'],
    ['编程语言', 'C#'],
    ['源程序量', '约3155行（15个C#源文件）'],
    ['开发目的', '为休闲益智类游戏爱好者提供宠物主题小游戏...'],
    ['面向领域/行业', '娱乐游戏 / 移动游戏 / 休闲益智'],
    ['主要功能', '关卡挑战、宠物系统、铲屎官成长、货币系统、局外养成...'],
    ['技术特点', '团结引擎、数据驱动、FSM状态机、BFS死局检测...'],
]

table = doc.add_table(rows=len(table_data), cols=2)
table.style = 'Table Grid'
for i, row_data in enumerate(table_data):
    row = table.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

out_path = 'F:/WorkBuddy/Claw/Design_Test/疯狂铲屎官_软著申请表填写内容.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
