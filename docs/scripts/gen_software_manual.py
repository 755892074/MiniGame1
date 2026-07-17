# -*- coding: utf-8 -*-
"""
软著申请 — 软件说明书生成器
基于游戏截图和功能描述，生成标准格式的 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = r"F:\WorkBuddy\Claw\Design_Test\疯狂铲屎官_软件说明书.docx"
SCREENSHOT_PATH = r"F:\WorkBuddy\H5MiniGame\MiniGame1\MiniGame1_Project\doc\游戏截图\6a8a00e6-00f7-4b66-bf16-f98ce37600ce.png"
GAME_NAME = "疯狂铲屎官"
VERSION = "V1.0"

def set_run_font(run, font_name='宋体', font_size=Pt(12), bold=False):
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_heading_para(doc, text, level=1):
    para = doc.add_paragraph()
    if level == 1:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(16), bold=True)
    elif level == 2:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(14), bold=True)
    else:
        run = para.add_run(text)
        set_run_font(run, font_name='黑体', font_size=Pt(12), bold=True)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_body_para(doc, text, indent=0, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_name='宋体', font_size=Pt(12), bold=bold)
    para.paragraph_format.first_line_indent = Cm(indent)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    return para

def main():
    doc = Document()
    
    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # ===== 封面 =====
    doc.add_paragraph()  # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{GAME_NAME}")
    set_run_font(run, font_name='黑体', font_size=Pt(26), bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("软件说明书")
    set_run_font(run, font_name='黑体', font_size=Pt(22), bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"版本号：{VERSION}\n日期：2026年07月")
    set_run_font(run, font_name='宋体', font_size=Pt(14))
    
    doc.add_page_break()
    
    # ===== 一、软件概述 =====
    add_heading_para(doc, "一、软件概述", 1)
    add_body_para(doc, f"《{GAME_NAME}》是一款基于Unity引擎开发的休闲益智类小游戏。游戏以宠物喂食为核心玩法，玩家需要通过点击操作将不同种类的食物从碗中倒出并重新分配，使相同食物填满同一个碗，从而触发喂食动作，按顺序满足排队宠物的需求，获得分数并通关。")
    add_body_para(doc, "游戏融合了经典的倒水排序（Water Sort）玩法与宠物题材包装，操作简单直观，适合各年龄段玩家。游戏共设计了50个精心编排的关卡，难度由浅入深，并配有撤销、加碗、打乱等辅助功能，提升游戏体验。")
    
    # ===== 二、运行环境 =====
    add_heading_para(doc, "二、运行环境", 1)
    add_body_para(doc, "开发环境：Unity 2022.3.62t1（团结引擎）")
    add_body_para(doc, "开发语言：C#")
    add_body_para(doc, "目标平台：WebGL（微信小游戏、抖音小游戏）")
    add_body_para(doc, "运行要求：支持HTML5的浏览器，屏幕分辨率建议750×1334及以上")
    
    # ===== 三、功能模块描述 =====
    add_heading_para(doc, "三、功能模块描述", 1)
    
    add_heading_para(doc, "3.1 游戏入口与关卡选择", 2)
    add_body_para(doc, "游戏启动后进入关卡选择界面。界面上方显示游戏名称《疯狂铲屎官》，下方以网格形式展示已解锁的关卡。玩家点击任意关卡即可进入游戏。")
    add_body_para(doc, "关卡共50关，分为简单、中等、困难三个难度等级。已通关的关卡会显示星级评价（1~3星），未解锁的关卡显示为锁定状态。")
    
    add_heading_para(doc, "3.2 核心玩法模块", 2)
    add_body_para(doc, "游戏主界面分为三个区域：")
    add_body_para(doc, "（1）顶部信息栏：显示当前关卡名称（如“第1关”）、当前得分/目标得分（如“得分：0/110”）、已用步数（“步数：0”）。")
    add_body_para(doc, "（2）宠物队列区：位于界面上方，显示当前等待喂食的宠物头像及名称。队首宠物是当前最优先需要被喂食的对象，成功喂食后该宠物离开队列，下一个宠物上前。")
    add_body_para(doc, "（3）碗区域：位于界面中下部，显示多个盛有不同食物的碗。玩家点击一个碗选取其中顶层相同食物，再点击另一个碗进行倾倒。只有当目标碗为空，或目标碗顶层食物与手上食物相同时，倾倒才能成功。")
    add_body_para(doc, "当某个碗被同一种食物完全填满时，该碗触发“完成”状态，对应宠物的食物需求被满足，宠物获得喂食，玩家根据宠物在队列中的位置获得不同分数（队首100分、第二位60分、第三位及以后30分），并计入连击加分。")
    
    add_heading_para(doc, "3.3 操作功能", 2)
    add_body_para(doc, "撤销（↩）：撤销上一步操作，将碗和分数恢复到上一步状态。每关不限次数。")
    add_body_para(doc, "加碗（🥣+）：在当前关卡中增加一个空碗，提供更多操作空间。可用于突破死局。")
    add_body_para(doc, "打乱（🔀）：将当前选中碗内的食物随机打乱顺序，改变局面。每关限制使用次数。")
    add_body_para(doc, "重开（🔄）：重新开始当前关卡，碗内食物和宠物队列恢复初始状态，分数和步数清零。")
    add_body_para(doc, "选关：返回关卡选择界面，可切换至其他关卡。")
    
    add_heading_para(doc, "3.4 计分与星级系统", 2)
    add_body_para(doc, "计分规则：")
    add_body_para(doc, "· 队首宠物匹配：100分 + 连击加成（每连击+20分）")
    add_body_para(doc, "· 第二位宠物匹配：60分")
    add_body_para(doc, "· 第三位及以后：30分")
    add_body_para(doc, "星级判定：")
    add_body_para(doc, "· 3星：得分达到目标分的100%及以上")
    add_body_para(doc, "· 2星：得分达到目标分的70%~99%")
    add_body_para(doc, "· 1星：得分低于目标分的70%但仍通关")
    
    add_heading_para(doc, "3.5 死局检测", 2)
    add_body_para(doc, "系统内置轻量级BFS死局检测算法。当玩家操作导致当前局面无法完成所有宠物喂食时，系统会提示“死局”，并提供加碗救援选项。此机制确保玩家不会因操作失误而陷入无法继续的困境，提升游戏体验。")
    
    add_heading_para(doc, "3.6 关卡生成与难度曲线", 2)
    add_body_para(doc, "游戏采用离线批量生成关卡的方式。每个关卡配置包括：宠物种类数、碗容量、额外空碗数、最少解步数范围等参数。关卡生成算法通过随机分配食物 + BFS验证可解性 + 筛选最少步数的方式，确保每关都有解且难度可控。")
    add_body_para(doc, "50关的难度曲线设计遵循“3易1难”的波浪式上升规律，从3种宠物、碗容量3的新手关，逐步过渡到6种宠物、碗容量5的高难关，满足玩家成长体验。")
    
    # ===== 四、界面说明 =====
    add_heading_para(doc, "四、界面说明", 1)
    add_body_para(doc, "以下为游戏主界面截图，展示了核心玩法的各个元素：")
    
    # 插入截图
    doc.add_picture(SCREENSHOT_PATH, width=Cm(12))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("图1 游戏主界面（第1关）")
    set_run_font(run, font_size=Pt(10))
    
    add_body_para(doc, "界面元素说明：")
    add_body_para(doc, "A — 顶部信息栏：显示当前关卡、得分/目标分、步数")
    add_body_para(doc, "B — 操作按钮行：撤销、加碗、选关、重开")
    add_body_para(doc, "C — 宠物队列：显示等待喂食的宠物（队首优先）")
    add_body_para(doc, "D — 碗区域：显示盛有不同食物的碗，玩家在此进行倾倒操作")
    add_body_para(doc, "E — 食物图标：每个碗中的食物以图标形式分层显示，顶层食物决定可倾倒目标")
    
    # ===== 五、技术特点 =====
    add_heading_para(doc, "五、技术特点", 1)
    add_body_para(doc, "1. FSM状态机驱动：游戏采用有限状态机（FSM）管理游戏流程，包括Idle（待机）、Selected（已选碗）、Pouring（倾倒中）、Feeding（喂食中）、Win（通关）、Fail（失败）六种状态，确保状态流转清晰可控。")
    add_body_para(doc, "2. BFS关卡验证：关卡生成阶段使用广度优先搜索（BFS）算法验证关卡可解性，并计算最少解步数，确保每关都有解。运行时采用轻量BFS检测死局，限制搜索深度12层、状态数2000个，兼顾性能与准确性。")
    add_body_para(doc, "3. 事件驱动架构：游戏内各模块通过UnityEvent进行解耦通信，UI层与逻辑层分离，便于后续功能扩展和维护。")
    add_body_para(doc, "4. 数据驱动设计：关卡配置采用ScriptableObject存储，支持运行时加载。CSV配置表 + 编辑器工具实现离线批量生成，方便设计师调整难度参数。")
    add_body_para(doc, "5. 撤销与历史记录：系统维护操作历史栈，支持无限次撤销，记录内容包括源碗/目标碗、食物数量、分数、连击数、完成状态等，确保撤销操作精确还原。")
    
    # ===== 六、使用说明 =====
    add_heading_para(doc, "六、使用说明", 1)
    add_body_para(doc, "1. 启动游戏：打开浏览器，访问游戏链接即可进入。首次进入会显示关卡选择界面。")
    add_body_para(doc, "2. 选择关卡：点击已解锁的关卡进入游戏。新手建议从第1关开始。")
    add_body_para(doc, "3. 操作方法：")
    add_body_para(doc, "   · 点击碗：选中该碗，取出顶层相同食物")
    add_body_para(doc, "   · 点击另一碗：将食物倒入目标碗")
    add_body_para(doc, "   · 填满一个碗：触发喂食，对应宠物获得食物")
    add_body_para(doc, "4. 通关条件：所有宠物都被成功喂食，或得分达到目标分。")
    add_body_para(doc, "5. 失败条件：当前步数超过最大限制（部分关卡），或进入死局且无救援手段。")
    add_body_para(doc, "6. 辅助功能：善用撤销、加碗、打乱功能，帮助突破难点。")
    
    # ===== 七、版权说明 =====
    add_heading_para(doc, "七、版权说明", 1)
    add_body_para(doc, f"本软件《{GAME_NAME}》为原创作品，著作权归开发者所有。未经许可，任何单位和个人不得复制、传播、修改或用于商业用途。")
    add_body_para(doc, "软件中包含的美术素材（宠物形象、食物图标、碗等）均为原创或合法授权使用。")
    add_body_para(doc, "本软件运行于WebGL平台，支持微信小游戏和抖音小游戏等渠道发布。")
    
    # 保存
    doc.save(OUTPUT_PATH)
    file_size = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"软件说明书已生成: {OUTPUT_PATH}")
    print(f"文件大小: {file_size:.1f} KB")
    print("完成!")

if __name__ == '__main__':
    main()
